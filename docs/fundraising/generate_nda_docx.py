#!/usr/bin/env python3
"""Generate a professional bilateral NDA as a .docx file.

Usage:
    python generate_nda_docx.py \
        --party1 "Bridgepoint Intelligence Inc." \
        --party2 "John Smith" \
        --date "2026-04-15" \
        --output nda_john_smith.docx

Requirements:
    pip install python-docx
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print(
        "Error: python-docx is required. Install with: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


# ---------------------------------------------------------------------------
# NDA clause content
# ---------------------------------------------------------------------------

CLAUSES: list[tuple[str, str, list[str]]] = [
    (
        "1. DEFINITION OF CONFIDENTIAL INFORMATION",
        (
            '"Confidential Information" means any and all information disclosed '
            "by one Party (the \"Disclosing Party\") to the other Party (the "
            '"Receiving Party"), whether orally, in writing, electronically, or '
            "by any other means, that is designated as confidential or that "
            "reasonably should be understood to be confidential given the nature "
            "of the information and the circumstances of disclosure. "
            "Confidential Information includes, without limitation:"
        ),
        [
            (
                "Technical Information: software code, algorithms, machine learning "
                "models, system architecture, data structures, APIs, patent applications "
                "(filed or unfiled), patent specifications, trade secrets, know-how, "
                "inventions, research data, and technical documentation;"
            ),
            (
                "Business Information: business plans, financial projections, revenue "
                "models, pricing strategies, customer lists, partnership strategies, "
                "market analyses, competitive intelligence, and go-to-market plans;"
            ),
            (
                "Financial Information: financial statements, fundraising plans, "
                "capitalization tables, investor lists, term sheets, and valuation "
                "analyses; and"
            ),
            (
                "Patent and Intellectual Property Information: patent claims, prosecution "
                "strategies, prior art analyses, IP portfolio strategies, and licensing "
                "terms."
            ),
        ],
    ),
    (
        "2. OBLIGATIONS OF CONFIDENTIALITY",
        "The Receiving Party shall:",
        [
            (
                "hold the Confidential Information in strict confidence and protect it "
                "with at least the same degree of care that the Receiving Party uses to "
                "protect its own confidential information of like kind, but in no event "
                "less than reasonable care;"
            ),
            (
                "use the Confidential Information solely for the Purpose of evaluating "
                "a potential business relationship and for no other purpose;"
            ),
            (
                "not disclose the Confidential Information to any third party without the "
                "prior written consent of the Disclosing Party, except to Representatives "
                "who have a need to know and are bound by confidentiality obligations no "
                "less restrictive than those herein; and"
            ),
            (
                "limit access to the Confidential Information to employees, officers, "
                "directors, advisors, and professional consultants who have a need to "
                "know for the Purpose."
            ),
        ],
    ),
    (
        "3. TERM",
        (
            "This Agreement shall remain in effect for a period of two (2) years from "
            "the Effective Date, unless terminated earlier by either Party upon thirty "
            "(30) days' written notice. The obligations of confidentiality shall survive "
            "termination for two (2) years, except with respect to trade secrets, which "
            "shall be protected for as long as they remain trade secrets under applicable "
            "law."
        ),
        [],
    ),
    (
        "4. RETURN AND DESTRUCTION OF MATERIALS",
        (
            "Upon expiration or termination of this Agreement, or upon written request "
            "of the Disclosing Party, the Receiving Party shall promptly return or "
            "destroy all materials containing Confidential Information and provide "
            "written certification of such return or destruction within fifteen (15) "
            "days. The Receiving Party may retain one archival copy solely for "
            "compliance with legal or regulatory obligations, subject to the "
            "confidentiality obligations herein."
        ),
        [],
    ),
    (
        "5. NO LICENSE TO INTELLECTUAL PROPERTY",
        (
            "Nothing in this Agreement grants the Receiving Party any licence, right, "
            "title, or interest in or to any intellectual property of the Disclosing "
            "Party, including without limitation any patents, patent applications, "
            "copyrights, trademarks, trade secrets, or other proprietary rights. All "
            "Confidential Information remains the exclusive property of the Disclosing "
            "Party."
        ),
        [],
    ),
    (
        "6. NO OBLIGATION TO TRANSACT",
        (
            "This Agreement does not obligate either Party to enter into any further "
            "agreement, transaction, or business relationship. Either Party may decline "
            "to pursue discussions at any time and for any reason, without liability. "
            "Each Party acknowledges that the other may develop similar or competitive "
            "products independently, provided such development does not use the "
            "Disclosing Party's Confidential Information."
        ),
        [],
    ),
    (
        "7. REMEDIES",
        (
            "Each Party acknowledges that a breach may cause irreparable harm for which "
            "monetary damages may be inadequate. The Disclosing Party shall be entitled "
            "to seek injunctive or equitable relief without proving actual damages or "
            "posting a bond. The substantially prevailing Party in any legal action "
            "shall be entitled to recover reasonable legal costs and attorney's fees."
        ),
        [],
    ),
    (
        "8. GOVERNING LAW AND JURISDICTION",
        (
            "This Agreement shall be governed by and construed in accordance with the "
            "laws of the Province of British Columbia and the federal laws of Canada "
            "applicable therein, without regard to conflict of laws principles. Each "
            "Party irrevocably submits to the exclusive jurisdiction of the courts of "
            "the Province of British Columbia."
        ),
        [],
    ),
    (
        "9. NON-SOLICITATION",
        (
            "During the term and for twelve (12) months following termination, neither "
            "Party shall directly or indirectly solicit for employment any employee, "
            "contractor, or consultant of the other Party with whom the soliciting Party "
            "had contact in connection with the Purpose. This restriction does not apply "
            "to general solicitations not specifically directed at the other Party's "
            "personnel, or to persons who initiate contact independently."
        ),
        [],
    ),
    (
        "10. RESIDUAL KNOWLEDGE",
        (
            'Nothing herein prevents either Party from using "Residual Knowledge" — '
            "ideas, concepts, know-how, or techniques retained in the unaided memory "
            "of Representatives who had access to Confidential Information, without "
            "reference to any written or electronic record. This does not grant a "
            "licence under any intellectual property right, permit disclosure to third "
            "parties, or allow intentional memorization for exploitation purposes."
        ),
        [],
    ),
]


def build_document(party1: str, party2: str, date: str) -> Document:
    """Build and return a python-docx Document containing the bilateral NDA."""
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # -- Styles --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # -- Title --
    title = doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"

    # -- Parties and date --
    doc.add_paragraph(
        f"This Mutual Non-Disclosure Agreement (the \"Agreement\") is entered into "
        f"as of {date} (the \"Effective Date\")"
    )

    doc.add_paragraph("BETWEEN:")
    p = doc.add_paragraph()
    p.add_run(f"(1)  {party1}").bold = True
    p.add_run('  (the "First Party")')

    doc.add_paragraph("AND:")
    p = doc.add_paragraph()
    p.add_run(f"(2)  {party2}").bold = True
    p.add_run('  (the "Second Party")')

    doc.add_paragraph(
        '(each a "Party" and collectively the "Parties")'
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -- Recitals --
    doc.add_heading("RECITALS", level=1)
    doc.add_paragraph(
        "WHEREAS the Parties wish to explore a potential business relationship "
        "relating to investment in and/or commercial engagement with technology "
        'for cross-border payment processing (the "Purpose"); and'
    )
    doc.add_paragraph(
        "WHEREAS in connection with the Purpose, each Party may disclose to the "
        "other certain confidential and proprietary information;"
    )
    doc.add_paragraph(
        "NOW THEREFORE, in consideration of the mutual covenants herein contained, "
        "and for other good and valuable consideration (the receipt and sufficiency "
        "of which are hereby acknowledged), the Parties agree as follows:"
    )

    # -- Clauses --
    for heading, body, sub_items in CLAUSES:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)
        for item in sub_items:
            p = doc.add_paragraph(item, style="List Bullet")

    # -- General provisions --
    doc.add_heading("11. GENERAL PROVISIONS", level=2)
    general = [
        (
            "Entire Agreement. ",
            "This Agreement constitutes the entire agreement between the Parties "
            "with respect to the subject matter hereof."
        ),
        (
            "Amendment. ",
            "This Agreement may only be amended by a written instrument signed by "
            "both Parties."
        ),
        (
            "Severability. ",
            "If any provision is held invalid, the remaining provisions continue "
            "in full force."
        ),
        (
            "Assignment. ",
            "Neither Party may assign this Agreement without prior written consent, "
            "except to a successor in a merger or acquisition."
        ),
        (
            "Counterparts. ",
            "This Agreement may be executed in counterparts, including by "
            "electronic signature."
        ),
    ]
    for label, text in general:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(text)

    # -- Signature blocks --
    doc.add_page_break()
    doc.add_heading("SIGNATURE PAGE", level=1)

    doc.add_paragraph(
        "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the "
        "Effective Date."
    )

    for party_label, party_name in [
        ("FIRST PARTY", party1),
        ("SECOND PARTY", party2),
    ]:
        doc.add_paragraph("")  # spacing
        p = doc.add_paragraph()
        p.add_run(party_label).bold = True

        if party_name:
            doc.add_paragraph(party_name)

        fields = ["Signature", "Name (Print)", "Title", "Email", "Address", "Date"]
        for field in fields:
            p = doc.add_paragraph()
            p.add_run(f"{field}: ").bold = True
            p.add_run("_" * 50)

    return doc


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a bilateral NDA as a professional .docx document."
    )
    parser.add_argument(
        "--party1",
        required=True,
        help="Name of the first party (e.g., 'Bridgepoint Intelligence Inc.')",
    )
    parser.add_argument(
        "--party2",
        required=True,
        help="Name of the second party (e.g., 'John Smith')",
    )
    parser.add_argument(
        "--date",
        required=True,
        help="Effective date of the NDA (e.g., '2026-04-15')",
    )
    parser.add_argument(
        "--output",
        required=True,
        help="Output .docx file path (e.g., 'nda_john_smith.docx')",
    )
    args = parser.parse_args()

    doc = build_document(args.party1, args.party2, args.date)

    output_path = Path(args.output)
    doc.save(str(output_path))
    print(f"NDA generated: {output_path.resolve()}")


if __name__ == "__main__":
    main()
